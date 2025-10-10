#!/usr/bin/env python3
from __future__ import annotations

import re
import sys
from pathlib import Path
from typing import List

try:
    from docx import Document
    from docx.enum.text import WD_BREAK
except Exception as exc:
    sys.stderr.write("python-docx is required. Install with: pip install python-docx\n")
    raise


TOKEN_PATTERN = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)")


def add_inline_markdown(paragraph, text: str) -> None:
    """Render inline markdown tokens (**bold**, *italic*, `code`) into a paragraph."""
    pos = 0
    for match in TOKEN_PATTERN.finditer(text):
        # text before token
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
        token = match.group(0)
        content = token[2:-2] if token.startswith("**") else token[1:-1]
        if token.startswith("**"):
            run = paragraph.add_run(content)
            run.bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(content)
            # monospace for code-like text
            run.font.name = "Courier New"
        else:  # *italic*
            run = paragraph.add_run(content)
            run.italic = True
        pos = match.end()
    # tail text
    if pos < len(text):
        paragraph.add_run(text[pos:])


def add_markdown_line(doc, line: str) -> None:
    s = line.rstrip("\n")
    if not s.strip():
        doc.add_paragraph("")
        return

    # horizontal separators
    if set(s.strip()) <= {"-"} and len(s.strip()) >= 3:
        # simple separator as empty paragraph
        p = doc.add_paragraph("")
        p.add_run("")
        return

    # headings
    if s.startswith("# "):
        doc.add_heading(s[2:].strip(), level=1)
        return
    if s.startswith("## "):
        doc.add_heading(s[3:].strip(), level=2)
        return
    if s.startswith("### "):
        doc.add_heading(s[4:].strip(), level=3)
        return
    if s.startswith("#### "):
        doc.add_heading(s[5:].strip(), level=4)
        return

    # bullet list
    if s.lstrip().startswith("- "):
        text = s.lstrip()[2:].strip()
        p = doc.add_paragraph(style="List Bullet")
        add_inline_markdown(p, text)
        return

    # default paragraph with inline formatting
    p = doc.add_paragraph("")
    add_inline_markdown(p, s)


def convert_md_to_docx(md_path: Path, docx_path: Path) -> None:
    lines: List[str] = md_path.read_text(encoding="utf-8").splitlines(True)
    doc = Document()
    # Title from first heading if exists
    title = None
    for ln in lines:
        if ln.startswith("# ") or ln.startswith("## "):
            title = ln.lstrip("# ").strip()
            break
    if title:
        doc.core_properties.title = title
    for ln in lines:
        add_markdown_line(doc, ln)
    doc.save(str(docx_path))


def main(argv: List[str]) -> int:
    if len(argv) < 3:
        sys.stderr.write("Usage: md_to_docx.py INPUT.md OUTPUT.docx\n")
        return 2
    md_path = Path(argv[1]).resolve()
    out_path = Path(argv[2]).resolve()
    convert_md_to_docx(md_path, out_path)
    print(f"Wrote {out_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))


